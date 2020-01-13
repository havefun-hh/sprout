import docx
import re
import pandas as pd
from Colored import Colored_1
import os
import glob


def replace_affix(doc, old, new):
    list = []
    for i in range(len(doc)):
        result = re.sub(f'\B{old}\W', f'{new}' + ' ', doc[i])
        list.append(result)
    return list

def replace_space(doc):
    list = []
    for i in range(len(doc)):
        result = re.sub('\B. \B', '.\n', doc[i])
        list.append(result)
    return list

def count(doc, old):
    num = 0
    for i in range(len(doc)):
        num += len(re.findall(f'\B{old}\W', doc[i]))
    print(f'{old}个数：{num}')
    return num

def merge(doc):
    n = len(doc)
    i = 0
    merge = ''
    while i < n:
        merge = merge + doc[i] + '\n'
        i += 1
    return merge

"""==================================================================================================================================="""
inpath = r"D:/Anaconda files/word文件处理/中文"
uipath = str(inpath)
folderlist = os.listdir(uipath)

# outer_path = r"D:/Anaconda files/word文件处理/computer science"         #修改双引号内路径
# folderlist = os.listdir(outer_path)
read_files = []
for i in range(len(folderlist)):
    for j in range(2015, 2020, 1):
        read_files.append(glob.glob('D:/Anaconda files/word文件处理/中文/' + folderlist[i] + '/' + str(j) + '/'+'*'+'.docx'))
all_file_names = [j for i in read_files for j in i]

df1 = pd.DataFrame(columns=['ion', 'ions', 'ment', 'ments', 'ness', 'nesses', 'ity', 'ities'])
for i in range(len(all_file_names)):
    file = docx.Document(all_file_names[i])
    doc = []
    for j in range(len(file.paragraphs)):
        doc.append(file.paragraphs[j].text)
    color = Colored_1()
    doc = replace_space(doc)
    doc1 = replace_affix(doc, 'ion', color.red_yellow('ion'))           #替换词缀
    doc2 = replace_affix(doc1, 'ions', color.red_yellow('ions'))
    doc3 = replace_affix(doc2, 'ment', color.red_yellow('ment'))
    doc4 = replace_affix(doc3, 'ments', color.red_yellow('ments'))
    doc5 = replace_affix(doc4, 'ness', color.red_yellow('ness'))
    doc6 = replace_affix(doc5, 'nesses', color.red_yellow('nesses'))
    doc7 = replace_affix(doc6, 'ity', color.red_yellow('ity'))
    doc8 = replace_affix(doc7, 'ities', color.red_yellow('ities'))
    print(merge(doc8))
    print("段落数:"+str(len(file.paragraphs)))
    data1 = []
    num1 = data1.append(count(doc, 'ion'))                              #计数
    num2 = data1.append(count(doc, 'ions'))
    num3 = data1.append(count(doc, 'ment'))
    num4 = data1.append(count(doc, 'ments'))
    num5 = data1.append(count(doc, 'ness'))
    num6 = data1.append(count(doc, 'nesses'))
    num7 = data1.append(count(doc, 'ity'))
    num8 = data1.append(count(doc, 'ities'))
    df1.loc[i + 1] = data1
df1.loc['Row_sum'] = df1.apply(lambda x: x.sum())

# #保存
# df1.to_csv('中文100_spyder.xlsx', sheet_name='Sheet1')
df1.to_csv('中文100_spyder.csv')




